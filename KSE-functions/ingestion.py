import logging
import io
import json

import azure.functions as func
from pypdf import PdfReader
from pypdf.errors import PyPdfError
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from langchain_core.documents import Document
from langchain_text_splitters import TokenTextSplitter

from shared import vector_store

bp = func.Blueprint()

CHUNK_SIZE_TOKENS = 512
CHUNK_OVERLAP_TOKENS = 75

# cl100k_base matches the embedding model's tokenizer
splitter = TokenTextSplitter(
    chunk_size=CHUNK_SIZE_TOKENS,
    chunk_overlap=CHUNK_OVERLAP_TOKENS,
    encoding_name="cl100k_base",
)


def get_doc_name_and_extension(blob_name: str) -> tuple[str, str]:
    filename = blob_name.split("/")[-1]
    if "." not in filename:
        raise ValueError(f"File has no extension, cannot determine type: {filename}")
    extension = filename.rsplit(".", 1)[-1].lower()
    return filename, extension


def make_chunk_id(doc_name: str, chunk_index: int) -> str:
    # deterministic ID = idempotent re-uploads overwrite instead of duplicating
    safe_name = doc_name.replace(".", "_").replace(" ", "_")
    return f"{safe_name}_{chunk_index}"


# hand-rolled extraction since LangChain loaders need file paths, not bytes

def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            pages_text.append(page_text)
    return "\n".join(pages_text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = DocxDocument(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in document.paragraphs)


def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")


def extract_text(file_bytes: bytes, extension: str) -> str:
    if extension == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif extension == "docx":
        return extract_text_from_docx(file_bytes)
    elif extension == "txt":
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{extension}")


def make_signal(target: str, doc_name: str, **extra) -> str:
    return json.dumps({
        "target": target,
        "arguments": [{"docName": doc_name, **extra}],
    })


@bp.generic_output_binding(
    arg_name="signalRMessages",
    type="signalR",
    hubName="ingestion",
    connectionStringSetting="AzureSignalRConnectionString",
)
@bp.blob_trigger(
    arg_name="myblob",
    path="documents/{name}",
    connection="AzureWebJobsStorage",
)
def blob_ingestion_trigger(myblob: func.InputStream, signalRMessages: func.Out[str]):
    logging.info(f"Blob trigger fired for: {myblob.name}")

    try:
        doc_name, extension = get_doc_name_and_extension(myblob.name)
        file_bytes = myblob.read()
        text = extract_text(file_bytes, extension)

    except ValueError as e:
        logging.warning(f"Skipping {myblob.name}: {e}")
        return

    except (PyPdfError, PackageNotFoundError) as e:
        logging.error(f"Failed to parse {myblob.name}: {e}")
        signalRMessages.set(make_signal("docFailed", myblob.name, error=str(e)))
        return

    if not text.strip():
        logging.warning(f"No extractable text in {doc_name}, skipping")
        signalRMessages.set(make_signal("docFailed", doc_name, error="No extractable text found"))
        return

    doc = Document(page_content=text, metadata={"doc_name": doc_name})
    chunks = splitter.split_documents([doc])
    chunk_ids = [make_chunk_id(doc_name, i) for i in range(len(chunks))]

    try:
        # embeds + upserts in one call
        vector_store.add_documents(documents=chunks, ids=chunk_ids)
    except Exception as e:
        logging.error(f"Failed to embed/upsert chunks for {doc_name}: {e}")
        signalRMessages.set(make_signal("docFailed", doc_name, error="Embedding or index write failed"))
        return

    logging.info(f"Upserted {len(chunks)} chunks for {doc_name}")
    signalRMessages.set(make_signal("docIndexed", doc_name, chunkCount=len(chunks)))