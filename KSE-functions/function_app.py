import logging
import io
import os
import json

import azure.functions as func
from azurefunctions.extensions.http.fastapi import Request, StreamingResponse

from pypdf import PdfReader
from pypdf.errors import PyPdfError
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from langchain_core.documents import Document
from langchain_core.messages import SystemMessage, HumanMessage
from langchain_text_splitters import TokenTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import AzureSearch

app = func.FunctionApp(http_auth_level=func.AuthLevel.ANONYMOUS)

# --- Config -------------------------------------------------------------

SEARCH_ENDPOINT = os.environ["AZURE_SEARCH_ENDPOINT"]
SEARCH_KEY = os.environ["AZURE_SEARCH_KEY"]
SEARCH_INDEX_NAME = "kse-chunks"

CHUNK_SIZE_TOKENS = 512
CHUNK_OVERLAP_TOKENS = 75
TOP_K_CHUNKS = 4

# TODO: replace with your Static Web App URL once deployed, keep localhost for local dev
CORS_HEADERS = {
    "Access-Control-Allow-Origin": "http://localhost:5173",
    "Access-Control-Allow-Methods": "POST, OPTIONS",
    "Access-Control-Allow-Headers": "Content-Type",
}

# --- Shared LangChain components (initialized once per warm instance) ---

embeddings = OpenAIEmbeddings(model="text-embedding-3-small", timeout=15)

splitter = TokenTextSplitter(
    chunk_size=CHUNK_SIZE_TOKENS,
    chunk_overlap=CHUNK_OVERLAP_TOKENS,
    encoding_name="cl100k_base",
)

vector_store = AzureSearch(
    azure_search_endpoint=SEARCH_ENDPOINT,
    azure_search_key=SEARCH_KEY,
    index_name=SEARCH_INDEX_NAME,
    embedding_function=embeddings.embed_query,
)

llm = ChatOpenAI(model="gpt-4o-mini", temperature=0, streaming=True, timeout=30)


# --- File name handling ---------------------------------------------------

def get_doc_name_and_extension(blob_name: str) -> tuple[str, str]:
    filename = blob_name.split("/")[-1]
    if "." not in filename:
        raise ValueError(f"File has no extension, cannot determine type: {filename}")
    extension = filename.rsplit(".", 1)[-1].lower()
    return filename, extension


def make_chunk_id(doc_name: str, chunk_index: int) -> str:
    safe_name = doc_name.replace(".", "_").replace(" ", "_")
    return f"{safe_name}_{chunk_index}"


# --- Text extraction ---------------------------------------------------------

def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            pages_text.append(page_text)
    return "\n".join(pages_text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = DocxDocument(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in document.paragraphs)


def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")


def extract_text(file_bytes: bytes, extension: str) -> str:
    if extension == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif extension == "docx":
        return extract_text_from_docx(file_bytes)
    elif extension == "txt":
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{extension}")


# --- Blob trigger: ingestion pipeline --------------------------------------

@app.blob_trigger(
    arg_name="myblob",
    path="documents/{name}",
    connection="AzureWebJobsStorage"
)
def blob_ingestion_trigger(myblob: func.InputStream):
    logging.info(f"Blob trigger fired for: {myblob.name}")

    try:
        doc_name, extension = get_doc_name_and_extension(myblob.name)
        file_bytes = myblob.read()
        text = extract_text(file_bytes, extension)

    except ValueError as e:
        logging.warning(f"Skipping {myblob.name}: {e}")
        return

    except (PyPdfError, PackageNotFoundError) as e:
        logging.error(f"Failed to parse {myblob.name}: {e}")
        return

    if not text.strip():
        logging.warning(f"No extractable text in {doc_name}, skipping")
        return

    doc = Document(page_content=text, metadata={"doc_name": doc_name})
    chunks = splitter.split_documents([doc])
    chunk_ids = [make_chunk_id(doc_name, i) for i in range(len(chunks))]

    try:
        vector_store.add_documents(documents=chunks, ids=chunk_ids)
    except Exception as e:
        logging.error(f"Failed to embed/upsert chunks for {doc_name}: {e}")
        return

    logging.info(f"Upserted {len(chunks)} chunks for {doc_name}")


# --- HTTP trigger: query pipeline -------------------------------------------

RAG_SYSTEM_PROMPT = (
    "You are a helpful assistant answering questions using only the provided context. "
    "If the context doesn't contain the answer, say you don't know — don't make one up."
)


async def generate_answer_stream(question: str):
    """Async generator: retrieve context, then stream the LLM's answer as SSE events."""
    logging.info(f"Query received: {question}")

    docs = vector_store.similarity_search(question, k=TOP_K_CHUNKS)
    logging.info(f"Retrieved {len(docs)} chunks")
    context = "\n\n---\n\n".join(d.page_content for d in docs)

    messages = [
        SystemMessage(content=RAG_SYSTEM_PROMPT),
        HumanMessage(content=f"Context:\n{context}\n\nQuestion: {question}"),
    ]

    async for chunk in llm.astream(messages):
        if chunk.content:
            payload = json.dumps({"token": chunk.content})
            yield f"data: {payload}\n\n"

    yield "data: [DONE]\n\n"


@app.route(route="query", methods=[func.HttpMethod.POST, func.HttpMethod.OPTIONS])
async def query_http_trigger(req: Request) -> StreamingResponse:
    if req.method == "OPTIONS":
        # Preflight check — no body needed, just the CORS headers
        async def empty():
            return
            yield  # makes this a generator that yields nothing

        return StreamingResponse(empty(), status_code=204, headers=CORS_HEADERS)

    body = await req.json()
    question = body.get("question")

    if not question:
        async def error_stream():
            yield f"data: {json.dumps({'error': 'Missing question field'})}\n\n"
        return StreamingResponse(
            error_stream(), media_type="text/event-stream", headers=CORS_HEADERS
        )

    return StreamingResponse(
        generate_answer_stream(question),
        media_type="text/event-stream",
        headers=CORS_HEADERS,
    )